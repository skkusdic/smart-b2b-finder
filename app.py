import json
import re as _re
import sys
from datetime import date as _date
from io import BytesIO as _BytesIO
import streamlit as st
import pandas as pd
import streamlit.components.v1 as components
from concurrent.futures import ThreadPoolExecutor, as_completed

sys.stdout.reconfigure(encoding="utf-8")

from pipeline.db import init_db, get_connection
from pipeline.filter import INDUSTRY_MAP, stage1_macro_filter
from pipeline.dart_collector import search_corp_by_name, fetch_corp_realtime
from agents.analyzer import (
    WEIGHT_PROFILES, compute_numeric_scores, get_weights,
    fetch_business_text, score_rag_single, _save_stage2_results,
)
from agents.report_writer import generate_report, DOC_TYPES, RECIPIENT_ROLES
from optimizer.milp import run_stage3

try:
    init_db()
except Exception:
    pass

st.set_page_config(
    page_title="Findly — Smart B2B Finder",
    layout="wide",
    initial_sidebar_state="collapsed",
)

RAG_LABEL = {
    "problem_fit":          "문제 일치도",
    "digital_willingness":  "디지털 전환 의지",
    "investment_direction": "신사업 투자 방향",
    "external_adoption":    "외부 솔루션 도입 이력",
    "decision_structure":   "의사결정 구조",
}

_RAG_TOOLTIP = {
    "problem_fit":          "우리 제품이 해당 기업의 실제 운영 문제와 얼마나 일치하는지. 공시·사업보고서 기반 Claude 판단.",
    "digital_willingness":  "IT·디지털 전환에 대한 기업의 적극성. R&D 투자, 스마트팩토리, DX 공시 등을 근거로 평가.",
    "investment_direction": "신사업·기술 투자 방향이 우리 솔루션 카테고리와 겹치는가. 중장기 사업 계획 공시 기반.",
    "external_adoption":    "과거 외부 소프트웨어·컨설팅 도입 이력. 경험이 많을수록 구매 장벽이 낮음.",
    "decision_structure":   "구매 결정 구조의 복잡도. 분권화된 조직일수록 담당자 접근이 용이.",
}

_SCORE_GUIDE = "점수 기준 — 7~10: 적합(녹색) · 4~6: 보통(주황) · 0~3: 주의(빨강) | Claude가 공시·사업보고서 원문을 분석해 항목별로 판단합니다."

# 산업군별 적합 제품 카테고리 매핑
INDUSTRY_TO_CATEGORIES = {
    "IT/소프트웨어":        ["SaaS/B2B 소프트웨어", "AI/데이터 솔루션", "컨설팅/서비스", "기타"],
    "제조업(전자/반도체)":  ["하드웨어/IoT", "AI/데이터 솔루션", "SaaS/B2B 소프트웨어", "기타"],
    "제조업(기계/자동차)":  ["하드웨어/IoT", "SaaS/B2B 소프트웨어", "컨설팅/서비스", "기타"],
    "제조업(화학/소재)":    ["SaaS/B2B 소프트웨어", "컨설팅/서비스", "기타"],
    "제조업(전체)":         ["하드웨어/IoT", "SaaS/B2B 소프트웨어", "AI/데이터 솔루션", "컨설팅/서비스", "기타"],
    "바이오/헬스케어":      ["AI/데이터 솔루션", "SaaS/B2B 소프트웨어", "컨설팅/서비스", "기타"],
    "금융":                 ["SaaS/B2B 소프트웨어", "AI/데이터 솔루션", "컨설팅/서비스", "기타"],
    "유통/물류":            ["SaaS/B2B 소프트웨어", "하드웨어/IoT", "AI/데이터 솔루션", "기타"],
    "건설/부동산":          ["SaaS/B2B 소프트웨어", "하드웨어/IoT", "컨설팅/서비스", "기타"],
    "서비스업":             ["SaaS/B2B 소프트웨어", "AI/데이터 솔루션", "컨설팅/서비스", "기타"],
    "전체":                 ["SaaS/B2B 소프트웨어", "AI/데이터 솔루션", "하드웨어/IoT", "컨설팅/서비스", "기타"],
}

for k, v in [("stage1_count", 0), ("stage2_result", None), ("excluded_count", 0),
             ("weights_used", {}), ("search_result", None), ("excluded_details", []),
             ("last_industry", "전체"), ("report_result", None), ("report_corp_name", ""),
             ("report_doc_type", ""), ("last_product_desc", ""), ("last_category", "기타"),
             ("profile_saved", False), ("show_manual_form", False), ("manual_corp_name", "")]:
    if k not in st.session_state:
        st.session_state[k] = v

try:
    conn = get_connection()
    total_in_db = conn.execute(
        "SELECT COUNT(*) FROM companies WHERE year=?", (2025,)
    ).fetchone()[0]
    conn.close()
except Exception:
    total_in_db = 0

r = st.session_state.stage2_result
avg_score_str = f"{r['total_score'].mean():.1f}" if (r is not None and len(r) > 0) else "—"

_CSS = """
* { font-family: 'Pretendard', 'Spoqa Han Sans Neo', 'Apple SD Gothic Neo', 'Noto Sans KR', sans-serif !important; }
#MainMenu, footer, header { display: none !important; }
[data-testid="stHeader"],
[data-testid="stToolbar"],
[data-testid="stDecoration"],
[data-testid="stStatusWidget"],
section[data-testid="stSidebar"],
[data-testid="collapsedControl"],
[data-testid="stSidebarCollapsedControl"] { display: none !important; }

html, body, .stApp {
    background:
        radial-gradient(circle 820px at 50% 42%,
            rgba(210, 40, 72, 0.22) 0%,
            rgba(232, 60, 90, 0.13) 42%,
            rgba(255, 180, 200, 0.07) 68%,
            transparent 100%
        ),
        #F5F5F7 !important;
}
.main,
.main .block-container,
div[class*="block-container"],
[data-testid="stAppViewBlockContainer"],
section[data-testid="stMain"] {
    padding: 0 !important;
    max-width: 100% !important;
}
.main { padding-top: 0 !important; margin-top: 0 !important; }
.stApp { padding-top: 0 !important; }
.element-container:first-of-type,
.stMarkdown:first-of-type { margin-top: 0 !important; padding-top: 0 !important; }

/* ── 다크 외부 컨테이너 ── */
[data-testid="stForm"] {
    background: #1a0a10 !important;
    border: 4px solid rgba(12,4,8,0.96) !important;
    border-radius: 24px !important;
    padding: 12px 16px 22px 16px !important;
    box-shadow: 0 14px 48px rgba(0,0,0,0.16) !important;
    margin-top: 0 !important;
    margin-bottom: 0 !important;
}

/* ── 흰색 플로팅 카드 (1단계 컬럼 그룹만) ── */
[data-testid="stForm"] [data-testid="stHorizontalBlock"] {
    background: #FFFFFF !important;
    border-radius: 16px !important;
    padding: 16px 18px 14px 18px !important;
    box-shadow: 0 8px 32px rgba(0,0,0,0.20) !important;
    margin-bottom: 4px !important;
}
/* ── 중첩 columns 는 투명하게 ── */
[data-testid="stForm"] [data-testid="stHorizontalBlock"] [data-testid="stHorizontalBlock"] {
    background: transparent !important;
    border-radius: 0 !important;
    padding: 0 !important;
    box-shadow: none !important;
    margin-bottom: 0 !important;
}

/* ── 인풋 & 셀렉트 ── */
[data-testid="stForm"] input, [data-testid="stForm"] textarea {
    background: rgba(0,0,0,0.04) !important;
    border: 1.5px solid rgba(0,0,0,0.13) !important;
    color: #1a0a10 !important;
    border-radius: 10px !important;
    font-size: 1.05rem !important;
}
[data-testid="stForm"] input::placeholder, [data-testid="stForm"] textarea::placeholder {
    color: rgba(26,10,16,0.33) !important;
}
[data-testid="stForm"] input:focus, [data-testid="stForm"] textarea:focus {
    border-color: rgba(232,48,90,0.5) !important;
    box-shadow: 0 0 0 3px rgba(232,48,90,0.11) !important;
}
[data-testid="stForm"] [data-baseweb="select"] > div {
    background: rgba(0,0,0,0.04) !important;
    border: 1.5px solid rgba(0,0,0,0.13) !important;
    color: #1a0a10 !important;
    border-radius: 10px !important;
}
[data-testid="stForm"] [data-baseweb="select"] svg { color: rgba(26,10,16,0.45) !important; }
[data-testid="stForm"] label { color: #3a1020 !important; font-size: 1.0rem !important; }
[data-testid="stForm"] [data-baseweb="slider"] [role="slider"] { background: #E8305A !important; }
[data-testid="stForm"] [data-baseweb="slider"] div { background: rgba(232,48,90,0.18) !important; }
/* 슬라이더 숫자 배경 제거 + 진한 색 */
[data-testid="stForm"] [data-testid="stSlider"] [data-baseweb="slider"] div[style*="background"] {
    background: transparent !important;
}
[data-testid="stForm"] [data-testid="stSlider"] p,
[data-testid="stForm"] [data-testid="stSlider"] span {
    color: #1a0a10 !important; font-weight: 700 !important;
    background: transparent !important; text-shadow: none !important;
}

/* ── 버튼 ── */
.stButton > button,
[data-testid="stForm"] .stButton > button,
[data-testid="stFormSubmitButton"] > button {
    background: linear-gradient(135deg, #C4183C 0%, #E8305A 60%, #FF6080 100%) !important;
    color: white !important; border: none !important; border-radius: 10px !important;
    font-weight: 600 !important; font-size: 1.1rem !important;
    padding: 0.5rem 1.2rem !important;
    box-shadow: 0 4px 18px rgba(232,48,90,0.35) !important;
    transition: all 0.18s ease !important;
}
.stButton > button:hover, [data-testid="stFormSubmitButton"] > button:hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 6px 24px rgba(232,48,90,0.48) !important;
}
/* ── 프로필 등록 버튼 — JS로 개별 처리 (profile_btn_muted 클래스) ── */
button.profile-btn-muted {
    background: rgba(50,18,28,0.55) !important;
    color: rgba(255,255,255,0.65) !important;
    font-size: 0.88rem !important; font-weight: 500 !important;
    padding: 0.38rem 0.9rem !important;
    box-shadow: 0 1px 6px rgba(0,0,0,0.18) !important;
    border: 1px solid rgba(255,255,255,0.14) !important;
}
button.profile-btn-muted:hover {
    background: rgba(70,25,40,0.7) !important;
    box-shadow: 0 2px 10px rgba(0,0,0,0.22) !important;
    transform: none !important;
}

/* ── 메트릭 ── */
[data-testid="metric-container"] {
    background: #FFFFFF !important; border: 1px solid rgba(232,48,90,0.1) !important;
    border-radius: 16px !important; padding: 16px 18px !important;
    box-shadow: 0 2px 14px rgba(232,48,90,0.06) !important;
}
[data-testid="metric-container"] label {
    color: #b07080 !important; font-size: 1.0rem !important;
    font-weight: 600 !important; text-transform: uppercase !important; letter-spacing: 0.8px !important;
}
[data-testid="metric-container"] [data-testid="metric-value"] {
    color: #E8305A !important; font-size: 2.0rem !important; font-weight: 700 !important;
}
.stProgress > div > div > div {
    background: linear-gradient(90deg, #E8305A, #FF8080) !important; border-radius: 8px !important;
}
.stAlert { border-radius: 12px !important; }
iframe { border: none !important; }

/* ── 탭 콘텐츠 좌우 여백 ── */
div[role="tabpanel"] {
    padding-left: 7.5rem !important;
    padding-right: 7.5rem !important;
}

/* ── 탭 pill 스타일 ── */
[data-baseweb="tab-list"] {
    background: #1a0a10 !important;
    border-radius: 28px !important;
    padding: 6px 8px !important;
    gap: 4px !important;
    border: 3px solid rgba(12,4,8,0.96) !important;
    box-shadow: 0 4px 20px rgba(0,0,0,0.22) !important;
    margin-bottom: 8px !important;
    margin-left: 7.5rem !important;
    width: fit-content !important;
}
[data-baseweb="tab"] {
    border-radius: 18px !important;
    color: rgba(255,255,255,0.5) !important;
    font-weight: 500 !important;
    background: transparent !important;
    border: none !important;
    padding: 9px 22px !important;
    font-size: 1.0rem !important;
    font-family: 'Pretendard', 'Apple SD Gothic Neo', 'Noto Sans KR', sans-serif !important;
    white-space: nowrap !important;
}
[data-baseweb="tab"]:hover {
    background: rgba(255,255,255,0.07) !important;
    color: rgba(255,255,255,0.75) !important;
}
[data-baseweb="tab"][aria-selected="true"] {
    background: white !important;
    color: #1a0a10 !important;
    font-weight: 700 !important;
    border-radius: 22px !important;
    box-shadow: 0 2px 10px rgba(0,0,0,0.12) !important;
}
[data-baseweb="tab-highlight"] { display: none !important; }
[data-baseweb="tab-border"]    { display: none !important; }
[data-testid="stTabs"] > div:first-child { overflow: visible !important; }

/* ── 보고서 패널: 컬럼 카드 ── */
[data-testid="stColumn"]:has(#rp-right-anchor) {
    background: #FFFFFF !important;
    border-radius: 18px !important;
    padding: 20px 24px 14px 24px !important;
    box-shadow: 0 2px 16px rgba(232,48,90,0.07) !important;
    border: 1px solid rgba(232,48,90,0.1) !important;
}
[data-testid="stColumn"]:has(#rp-left-anchor) {
    background: rgba(255,255,255,0.55) !important;
    border-radius: 18px !important;
    padding: 20px 20px 18px 20px !important;
    border: 1px solid rgba(232,48,90,0.08) !important;
}
/* 보고서 패널 인풋 */
[data-testid="stColumn"]:has(#rp-left-anchor) input,
[data-testid="stColumn"]:has(#rp-left-anchor) textarea {
    background: #fff !important;
    border: 1.5px solid rgba(0,0,0,0.13) !important;
    color: #1a0a10 !important;
    border-radius: 10px !important;
    font-size: 1.0rem !important;
}
[data-testid="stColumn"]:has(#rp-left-anchor) input:focus,
[data-testid="stColumn"]:has(#rp-left-anchor) textarea:focus {
    border-color: rgba(232,48,90,0.5) !important;
    box-shadow: 0 0 0 3px rgba(232,48,90,0.1) !important;
}
[data-testid="stColumn"]:has(#rp-left-anchor) [data-baseweb="select"] > div {
    background: #fff !important;
    border: 1.5px solid rgba(0,0,0,0.13) !important;
    color: #1a0a10 !important;
    border-radius: 10px !important;
}
[data-testid="stColumn"]:has(#rp-left-anchor) label { color: #3a1020 !important; }
[data-testid="stColumn"]:has(#rp-right-anchor) textarea {
    background: rgba(0,0,0,0.015) !important;
    border: 1px solid rgba(232,48,90,0.1) !important;
    border-radius: 12px !important;
    font-size: 0.95rem !important;
    line-height: 1.75 !important;
    color: #2a1018 !important;
}
/* 보고서 패널 공통 클래스 */
.rp-panel-header {
    background: linear-gradient(135deg,#1a0a10 0%,#2d0f1c 100%);
    border-radius: 12px; padding: 12px 16px;
    display: flex; align-items: center; gap: 8px;
    color: #fff; font-size: 1.0rem; font-weight: 700;
    margin-bottom: 12px; letter-spacing: 0.3px;
}
.rp-label {
    color: #3a1020; font-size: 0.855rem; font-weight: 600;
    margin: 10px 0 3px 0; display: block; letter-spacing: 0.15px;
}
.rp-tags-row { display: flex; flex-wrap: wrap; gap: 5px; margin: 4px 0 6px 0; }
.rp-tag {
    background: rgba(232,48,90,0.09); color: #C4183C;
    border: 1.5px solid rgba(232,48,90,0.22); border-radius: 20px;
    padding: 3px 11px; font-size: 0.82rem; font-weight: 500;
}
.rp-right-empty { padding: 56px 20px; text-align: center; }
.rp-right-empty-icon { font-size: 2.4rem; margin-bottom: 10px; }
.rp-right-empty-title { color: #b07080; font-size: 1.1rem; font-weight: 600; margin: 0 0 6px 0; }
.rp-right-empty-sub { color: #c0a0a8; font-size: 0.92rem; margin: 0; line-height: 1.65; }
.rp-doc-tabs { display: flex; gap: 4px; margin-bottom: 10px; }
.rp-doc-tab {
    padding: 5px 16px; border-radius: 20px;
    font-size: 0.86rem; font-weight: 500; color: rgba(26,10,16,0.38);
    border: 1.5px solid rgba(0,0,0,0.08);
}
.rp-doc-tab.active {
    background: #E8305A; color: #fff; border-color: #E8305A;
    box-shadow: 0 2px 8px rgba(232,48,90,0.28);
}
.rp-corp-label { color: rgba(26,10,16,0.38); font-size: 0.8rem; margin: 0 0 6px 2px; font-weight: 500; }
.rp-action-bar {
    display: flex; align-items: center;
    padding: 8px 0 4px 0; margin-top: 4px;
    border-top: 1px solid rgba(232,48,90,0.09);
}
.rp-action-hint { color: rgba(26,10,16,0.35); font-size: 0.78rem; }
"""

_FONT_SPOQA = "https://cdn.jsdelivr.net/gh/spoqa/spoqa-han-sans@latest/css/SpoqaHanSansNeo.css"
_FONT_PRETENDARD = "https://cdn.jsdelivr.net/gh/orioncactus/pretendard@v1.3.9/dist/web/static/pretendard.css"
components.html(f"""
<script>
(function(){{
    var pid='findly-styles';
    if(!window.parent.document.getElementById(pid)){{
        [{json.dumps(_FONT_SPOQA)},{json.dumps(_FONT_PRETENDARD)}].forEach(function(href){{
            var lk=window.parent.document.createElement('link');
            lk.rel='stylesheet'; lk.href=href;
            window.parent.document.head.appendChild(lk);
        }});
        var s=window.parent.document.createElement('style');
        s.id=pid; s.textContent={json.dumps(_CSS)};
        window.parent.document.head.appendChild(s);
    }}
    function markProfileBtn(){{
        var btns=window.parent.document.querySelectorAll('[data-testid="stFormSubmitButton"] > button');
        btns.forEach(function(b){{
            var t=(b.innerText||b.textContent||'').trim();
            if(t==='프로필 등록'){{
                b.classList.add('profile-btn-muted');
            }} else {{
                b.classList.remove('profile-btn-muted');
            }}
        }});
    }}
    [100,400,1000,2000].forEach(function(t){{setTimeout(markProfileBtn,t);}});
    var _obs=new MutationObserver(markProfileBtn);
    _obs.observe(window.parent.document.body,{{childList:true,subtree:true}});

    function fixTopGap(){{
        var sels=[
            '[data-testid="stMain"]',
            'section[data-testid="stMain"]',
            '.main',
            '.block-container',
            '[data-testid="stAppViewBlockContainer"]',
            '[data-testid="stAppViewContainer"]'
        ];
        sels.forEach(function(sel){{
            var el=window.parent.document.querySelector(sel);
            if(el){{
                el.style.setProperty('padding-top','0px','important');
                el.style.setProperty('margin-top','0px','important');
            }}
        }});
        var hdr=window.parent.document.querySelector('[data-testid="stHeader"]');
        if(hdr){{
            hdr.style.setProperty('display','none','important');
            hdr.style.setProperty('height','0','important');
        }}
    }}
    [0,80,300,800,2000].forEach(function(t){{setTimeout(fixTopGap,t);}});
    var _obs2=new MutationObserver(fixTopGap);
    _obs2.observe(window.parent.document.body,{{childList:true,subtree:true,attributes:true,attributeFilter:['style']}});
}})();
</script>
""", height=0, scrolling=False)

# ── Hero block (full-width white, no lateral padding) ─────────────────────────
_about_dd = (
    '<div style="position:absolute;z-index:9999;background:#FFFFFF;'
    'border:1px solid rgba(0,0,0,0.09);border-radius:14px;'
    'padding:16px 20px;box-shadow:0 8px 32px rgba(0,0,0,0.12);'
    'min-width:260px;margin-top:8px;right:0;">'
    '<p style="color:#1a0a10;font-weight:700;font-size:1.05rem;margin:0 0 10px 0;">findly</p>'
    '<p style="color:#888;font-size:0.82rem;margin:0 0 4px 0;"><b style="color:#555;">Author</b> — 이정원</p>'
    '<p style="color:#888;font-size:0.82rem;margin:0 0 4px 0;"><b style="color:#555;">APIs</b> — DART 공시정보 API, Anthropic Claude API</p>'
    '<p style="color:#888;font-size:0.82rem;margin:0 0 4px 0;"><b style="color:#555;">Stack</b> — Python · Streamlit · SQLite · PuLP(MILP)</p>'
    '<p style="color:#888;font-size:0.82rem;margin:0;"><b style="color:#555;">Model</b> — claude-haiku-4-5</p>'
    '</div>'
)
_target_svg = (
    '<svg width="240" height="240" viewBox="0 0 240 240" fill="none" xmlns="http://www.w3.org/2000/svg">'
    '<circle cx="120" cy="120" r="112" fill="rgba(232,48,90,0.05)" stroke="rgba(232,48,90,0.13)" stroke-width="1"/>'
    '<circle cx="120" cy="120" r="88" fill="rgba(232,48,90,0.09)" stroke="rgba(232,48,90,0.18)" stroke-width="1"/>'
    '<circle cx="120" cy="120" r="65" fill="rgba(232,48,90,0.14)" stroke="rgba(232,48,90,0.24)" stroke-width="1"/>'
    '<circle cx="120" cy="120" r="43" fill="rgba(232,48,90,0.55)" stroke="#C4183C" stroke-width="1.5"/>'
    '<circle cx="120" cy="120" r="22" fill="#C4183C"/>'
    '<circle cx="120" cy="120" r="9" fill="white" fill-opacity="0.55"/>'
    '<line x1="196" y1="44" x2="128" y2="120" stroke="#2a1020" stroke-width="7" stroke-linecap="round"/>'
    '<polygon points="128,120 142,110 134,126" fill="#1a0a0a"/>'
    '<rect x="192" y="36" width="14" height="5" rx="2.5" fill="#C4183C" transform="rotate(-45 196 44)"/>'
    '</svg>'
)
# ── Nav bar (흰색, 최상단 flush) ──────────────────────────────────────────────
_nav_html = (
    '<div style="background:#FFFFFF;padding:13px 2.5rem;'
    'display:flex;justify-content:space-between;align-items:center;'
    'border-bottom:1px solid rgba(0,0,0,0.06);">'
    '<div style="display:flex;align-items:center;gap:10px;">'
    '<div style="width:34px;height:34px;background:linear-gradient(135deg,#C4183C,#FF6080);'
    'border-radius:9px;display:flex;align-items:center;justify-content:center;'
    'box-shadow:0 3px 10px rgba(232,48,90,0.35);flex-shrink:0;">'
    '<span style="color:white;font-size:0.85rem;font-weight:800;font-family:\'Inter\',sans-serif;">f</span>'
    '</div>'
    '<span style="font-family:\'Inter\',sans-serif;font-size:1.85rem;font-weight:900;'
    'letter-spacing:-1px;background:linear-gradient(135deg,#C4183C 0%,#E8305A 45%,#FF6080 100%);'
    '-webkit-background-clip:text;-webkit-text-fill-color:transparent;'
    'background-clip:text;line-height:1;">findly</span>'
    '</div>'
    '<div style="display:flex;align-items:center;gap:22px;">'
    '<details style="cursor:pointer;position:relative;">'
    '<summary style="list-style:none;font-family:\'Inter\',sans-serif;font-size:0.95rem;'
    'font-weight:600;color:#999;outline:none;user-select:none;letter-spacing:0.2px;">About Us</summary>'
    + _about_dd +
    '</details>'
    f'<div style="background:rgba(232,48,90,0.07);border:1px solid rgba(232,48,90,0.14);'
    f'border-radius:20px;padding:4px 16px;white-space:nowrap;">'
    f'<span style="color:#b07080;font-size:0.9rem;font-weight:500;">DB </span>'
    f'<span style="color:#E8305A;font-size:1.15rem;font-weight:800;">{total_in_db:,}개</span>'
    f'</div>'
    '<div style="width:36px;height:36px;background:linear-gradient(135deg,#E8305A,#FF6080);'
    'border-radius:50%;display:flex;align-items:center;justify-content:center;'
    'font-size:0.82rem;color:white;font-weight:800;'
    'box-shadow:0 2px 10px rgba(232,48,90,0.32);flex-shrink:0;'
    'font-family:\'Inter\',sans-serif;">AI</div>'
    '</div>'
    '</div>'
)
st.markdown(_nav_html, unsafe_allow_html=True)

# ── Hero content (배경 박스 없음 — 페이지 그라디언트 위에 직접) ─────────────────
_hero_content_html = (
    '<div style="display:flex;align-items:center;justify-content:center;gap:60px;'
    'padding:44px 6rem 52px 6rem;">'
    '<div style="flex:1;max-width:540px;">'
    '<div style="display:inline-flex;align-items:center;gap:6px;'
    'background:rgba(232,48,90,0.08);border:1px solid rgba(232,48,90,0.22);'
    'border-radius:20px;padding:5px 15px;margin-bottom:18px;">'
    '<span style="color:#E8305A;font-size:0.85rem;">✦</span>'
    '<span style="color:#E8305A;font-size:0.85rem;font-weight:600;letter-spacing:0.3px;">AI 기반 B2B 타깃 분석 플랫폼</span>'
    '</div>'
    '<h1 style="font-size:2.55rem;font-weight:800;color:#1a0a10;'
    'font-family:\'Pretendard\',\'Apple SD Gothic Neo\',\'Noto Sans KR\',sans-serif;'
    'margin:0 0 11px 0;letter-spacing:-1.2px;line-height:1.2;">'
    'B2B 타깃, 데이터로 발굴하다</h1>'
    '<p style="color:#7a4050;font-size:1.1rem;margin:0;letter-spacing:0.2px;line-height:1.6;">'
    'DART 재무 데이터 · Claude RAG 분석 · MILP 최적화</p>'
    '</div>'
    '<div style="position:relative;width:360px;height:270px;flex-shrink:0;margin-left:1.5rem;">'
    '<div style="position:absolute;top:15px;left:60px;">' + _target_svg + '</div>'
    '<div style="position:absolute;top:8px;left:0;background:white;border-radius:14px;'
    'padding:10px 16px;box-shadow:0 4px 22px rgba(0,0,0,0.1);min-width:110px;">'
    '<p style="color:#aaa;font-size:0.68rem;font-weight:700;text-transform:uppercase;'
    'letter-spacing:0.8px;margin:0 0 4px 0;">유도 분석</p>'
    '<p style="color:#E8305A;font-size:1.6rem;font-weight:800;margin:0;line-height:1;">82%</p>'
    '</div>'
    '<div style="position:absolute;top:8px;right:0;background:white;border-radius:14px;'
    'padding:10px 16px;box-shadow:0 4px 22px rgba(0,0,0,0.1);min-width:120px;">'
    '<p style="color:#aaa;font-size:0.68rem;font-weight:700;text-transform:uppercase;'
    'letter-spacing:0.8px;margin:0 0 5px 0;">협력 적합도</p>'
    '<div style="display:flex;align-items:flex-end;gap:5px;">'
    '<p style="color:#1a0a10;font-size:1.6rem;font-weight:800;margin:0;line-height:1;">A</p>'
    '<div style="display:flex;align-items:flex-end;gap:2px;margin-bottom:4px;">'
    '<div style="width:4px;height:8px;background:#ffccd5;border-radius:2px;"></div>'
    '<div style="width:4px;height:13px;background:#ff8ca0;border-radius:2px;"></div>'
    '<div style="width:4px;height:20px;background:#E8305A;border-radius:2px;"></div>'
    '</div></div></div>'
    '<div style="position:absolute;bottom:8px;left:4px;background:white;border-radius:14px;'
    'padding:10px 16px;box-shadow:0 4px 22px rgba(0,0,0,0.1);min-width:130px;">'
    '<p style="color:#aaa;font-size:0.68rem;font-weight:700;text-transform:uppercase;'
    'letter-spacing:0.8px;margin:0 0 5px 0;">재무 안정성</p>'
    '<div style="display:flex;align-items:center;gap:7px;">'
    '<p style="color:#1a0a10;font-size:1.6rem;font-weight:800;margin:0;line-height:1;">상</p>'
    '<svg width="42" height="22" viewBox="0 0 42 22">'
    '<polyline points="0,18 10,14 20,9 30,5 42,1" fill="none" stroke="#E8305A"'
    ' stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round"/>'
    '</svg></div></div>'
    '</div>'
    '</div>'
)
st.markdown(_hero_content_html, unsafe_allow_html=True)

# ── Content (with lateral padding) ────────────────────────────────────────────
st.markdown("<div>", unsafe_allow_html=True)

tab_dashboard, tab_report, tab_history = st.tabs(["대시보드", "보고서 제작", "히스토리"])

# ── 모듈레벨 헬퍼 함수 ───────────────────────────────────────────────────────

def result_card(rank: int, row: pd.Series) -> str:
    total     = row["total_score"]
    c         = "#E8305A" if total >= 70 else "#FF8C42" if total >= 50 else "#9a9aaa"
    bg        = "rgba(232,48,90,0.02)" if total >= 70 else "rgba(255,140,66,0.02)" if total >= 50 else "#FAFAFA"
    _nq      = row["corp_name"].replace(" ", "+")
    info_url = f"https://search.naver.com/search.naver?query={_nq}+기업정보"

    detail         = ""
    rag_scores_list = []
    try:
        d = json.loads(row.get("rag_detail", "{}"))
        for key, label in RAG_LABEL.items():
            if key in d and isinstance(d[key], dict):
                sc     = d[key].get("score", 0)
                reason = d[key].get("reason", "")
                rag_scores_list.append((label, sc, reason))
                pct   = sc / 10 * 100
                # A: 점수 해석 배지
                if sc >= 7:
                    b_color, b_bg, b_text = "#1a7a40", "#e8f5ee", "적합"
                elif sc >= 4:
                    b_color, b_bg, b_text = "#b06000", "#fff3e0", "보통"
                else:
                    b_color, b_bg, b_text = "#C4183C", "#fff0f3", "주의"
                # C: 항목별 툴팁
                tooltip = _RAG_TOOLTIP.get(key, "")
                detail += (
                    f'<div style="padding:6px 0;border-bottom:1px solid rgba(232,48,90,0.06);">'
                    f'<div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:3px;">'
                    f'<div style="display:flex;align-items:center;gap:4px;">'
                    f'<span style="color:#6b3040;font-size:1.0rem;">{label}</span>'
                    f'<span title="{tooltip}" style="display:inline-flex;align-items:center;justify-content:center;'
                    f'width:15px;height:15px;border-radius:50%;background:rgba(232,48,90,0.1);'
                    f'color:#c07080;font-size:0.68rem;cursor:help;font-weight:700;flex-shrink:0;">?</span>'
                    f'</div>'
                    f'<div style="display:flex;align-items:center;gap:5px;">'
                    f'<span style="background:{b_bg};color:{b_color};font-size:0.72rem;font-weight:700;'
                    f'border-radius:4px;padding:1px 6px;">{b_text}</span>'
                    f'<span style="color:#E8305A;font-size:1.0rem;font-weight:600;">{sc}/10</span>'
                    f'</div></div>'
                    f'<div style="background:#FFF0F3;border-radius:4px;height:4px;overflow:hidden;margin-bottom:3px;">'
                    f'<div style="background:linear-gradient(90deg,#E8305A,#FF8080);width:{pct}%;height:100%;border-radius:4px;"></div>'
                    f'</div>'
                    f'<span style="color:#9a6070;font-size:0.92rem;">{reason}</span>'
                    f'</div>'
                )
    except Exception:
        pass

    # B: 이 기업을 선택한 이유 (상위 2개 RAG 항목 기반)
    selection_html = ""
    if rag_scores_list:
        top2 = sorted(rag_scores_list, key=lambda x: x[1], reverse=True)[:2]
        reason_str = " · ".join(f"{item[0]}({item[1]}/10)" for item in top2)
        selection_html = (
            f'<div style="background:linear-gradient(135deg,rgba(30,130,80,0.07),rgba(30,160,90,0.02));'
            f'border:1px solid rgba(30,160,80,0.15);border-radius:8px;padding:8px 12px;margin-bottom:10px;">'
            f'<p style="color:#1a6040;font-size:0.78rem;font-weight:700;text-transform:uppercase;'
            f'letter-spacing:0.8px;margin:0 0 3px 0;">이 기업을 선택한 이유</p>'
            f'<p style="color:#1a5030;font-size:0.92rem;margin:0;">'
            f'{reason_str} 점수가 높아 솔루션 수용 가능성이 큽니다.</p>'
            f'</div>'
        )

    # C: RAG 섹션 헤더 툴팁
    rag_header = (
        f'<div style="display:flex;align-items:center;gap:6px;margin:0 0 6px 0;">'
        f'<p style="color:#9a4a5a;font-size:0.82rem;font-weight:700;text-transform:uppercase;'
        f'letter-spacing:0.8px;margin:0;">RAG 항목별 상세</p>'
        f'<span title="{_SCORE_GUIDE}" style="display:inline-flex;align-items:center;justify-content:center;'
        f'width:16px;height:16px;border-radius:50%;background:rgba(232,48,90,0.1);'
        f'color:#c07080;font-size:0.7rem;cursor:help;font-weight:700;">?</span>'
        f'</div>'
    )

    return (
        f'<details style="background:{bg};border:1px solid rgba(232,48,90,0.09);'
        f'border-radius:13px;padding:12px 16px;margin-bottom:6px;cursor:pointer;">'
        f'<summary style="list-style:none;display:flex;align-items:center;gap:10px;outline:none;">'
        f'<span style="background:linear-gradient(135deg,#C4183C,#FF6080);color:white;'
        f'border-radius:7px;min-width:26px;height:26px;display:flex;align-items:center;'
        f'justify-content:center;font-size:1.0rem;font-weight:700;flex-shrink:0;">{rank}</span>'
        f'<span style="font-weight:600;font-size:1.15rem;color:#1a0a10;flex:1;">{row["corp_name"]}</span>'
        f'<span style="color:#c07080;font-size:0.95rem;">{row.get("market_type","-")} · {row.get("industry","-")}</span>'
        f'<a href="{info_url}" target="_blank" onclick="event.stopPropagation();" '
        f'style="color:#E8305A;font-size:0.78rem;text-decoration:none;border:1px solid rgba(232,48,90,0.3);'
        f'border-radius:6px;padding:2px 8px;white-space:nowrap;flex-shrink:0;">기업 정보 →</a>'
        f'<div style="display:flex;gap:5px;align-items:center;">'
        f'<span style="color:#c07080;font-size:0.95rem;">R&D {row["rd_score"]:.1f}</span>'
        f'<span style="color:#c07080;font-size:0.95rem;">부채 {row["debt_score"]:.1f}</span>'
        f'<span style="color:#c07080;font-size:0.95rem;">RAG {row["rag_score"]:.1f}</span>'
        f'<span style="color:{c};font-weight:700;font-size:1.2rem;background:rgba(232,48,90,0.07);'
        f'padding:2px 9px;border-radius:7px;">{total:.1f}</span>'
        f'</div></summary>'
        f'<div style="margin-top:10px;padding-top:10px;border-top:1px solid rgba(232,48,90,0.07);">'
        f'<p style="color:#c07080;font-size:0.95rem;margin:0 0 8px 0;">매출 {int(row.get("revenue") or 0):,}원</p>'
        f'{selection_html}'
        f'{rag_header}'
        f'{detail}</div></details>'
    )


def search_panel_html(sr: dict, profile: dict = None, stage2_df=None, total_db: int = 0) -> str:  # noqa: C901
    profile    = profile or {}
    p_industry = profile.get("industry", "")
    p_category = profile.get("category", "기타")
    p_desc     = profile.get("product_desc", "")
    p_desc_short = (p_desc[:80] + "…") if len(p_desc) > 80 else p_desc
    corp_name  = sr["name"]
    is_error   = sr.get("error", False)
    _src_tag   = sr.get("source", "DB")
    today      = _date.today().strftime("%Y.%m.%d")

    revenue_raw = sr.get("revenue", 0) or 0
    debt_ratio  = sr.get("debt_ratio", 0) or 0
    rev_str  = (f"{revenue_raw // 100_000_000:,}억 원" if revenue_raw >= 100_000_000
                else f"{revenue_raw:,}원" if revenue_raw else "정보 없음")
    debt_str = f"{debt_ratio:.0f}%" if debt_ratio else "정보 없음"
    if debt_ratio > 200:   debt_interp = "부채비율 매우 높음, 외부 지출 여력 제한적"
    elif debt_ratio > 100: debt_interp = "부채비율 높음, 신규 IT 투자 우선순위 낮을 수 있음"
    elif debt_ratio > 0:   debt_interp = "부채비율 안정적, 재무 측면 장애 적음"
    else:                  debt_interp = "부채비율 정보 없음"

    # RAG 데이터 파싱
    rag_d = {}
    try:
        rag_d = json.loads(sr.get("rag_detail", "{}"))
    except Exception:
        pass
    main_products  = rag_d.get("main_products", "")
    key_partners   = rag_d.get("key_partners", "")
    collab_signal  = rag_d.get("collaboration_signal", "")
    collab_reason  = rag_d.get("collaboration_reason", "")

    rag_items = [(key, RAG_LABEL[key], rag_d[key]["score"], rag_d[key].get("reason", ""))
                 for key in RAG_LABEL if key in rag_d and isinstance(rag_d[key], dict)]

    # ── 공통: 판정 결과 계산 ─────────────────────────────────────────────────
    if is_error:
        judgment_label = "협력 부적합"
        judgment_color = "#C4183C"
        judgment_bg    = "rgba(200,30,60,0.09)"
        total_display  = (sum(s for _, _, s, _ in rag_items) / len(rag_items) * 10) if rag_items else 0
    else:
        total_display = sr["total"]
        if total_display >= 70:
            judgment_label, judgment_color, judgment_bg = "협력 적합", "#1a7a40", "rgba(30,140,70,0.09)"
        elif total_display >= 50:
            judgment_label, judgment_color, judgment_bg = "협력 적합 (조건부)", "#b06000", "rgba(180,100,0,0.09)"
        else:
            judgment_label, judgment_color, judgment_bg = "협력 가능성 낮음", "#7a6000", "rgba(120,100,0,0.07)"

    pct         = min(total_display, 100)
    bar_color   = "#E8305A" if pct >= 70 else "#FF8C42" if pct >= 50 else "#9a9aaa"
    _nq         = corp_name.replace(" ", "+")
    dart_url    = f"https://search.naver.com/search.naver?query={_nq}+기업정보"

    # ── 순위 배지 (채택 케이스) ───────────────────────────────────────────────
    rank_badge = ""
    if not is_error and stage2_df is not None and len(stage2_df) > 0:
        match = stage2_df[stage2_df["corp_name"] == corp_name]
        if not match.empty:
            sdf      = stage2_df.sort_values("total_score", ascending=False).reset_index(drop=True)
            rank_pos = int(sdf[sdf["corp_name"] == corp_name].index[0]) + 1
            rank_of  = len(stage2_df)
            rank_badge = (
                f'<span style="background:rgba(232,48,90,0.1);color:#C4183C;'
                f'font-size:0.83rem;font-weight:700;border-radius:5px;padding:2px 10px;">'
                f'필터 탐색 {rank_pos}위 / {rank_of}개</span> '
            )
    if total_display >= 80:   pct_label = "상위 5%"
    elif total_display >= 70: pct_label = "상위 15%"
    elif total_display >= 60: pct_label = "상위 30%"
    elif total_display >= 50: pct_label = "상위 50%"
    else:                     pct_label = "평균 이하"

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 1 — 리포트 헤더
    # ════════════════════════════════════════════════════════════════════════
    S = '<div style="font-family:\'Pretendard\',\'Apple SD Gothic Neo\',\'Noto Sans KR\',sans-serif;color:#1a0a10;">'

    S += (
        f'<div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:16px;">'
        f'<span style="color:#aaa;font-size:0.9rem;">← 기업 분석 결과</span>'
        f'<a href="{dart_url}" target="_blank" style="color:#6a6a80;font-size:0.88rem;'
        f'text-decoration:none;border:1px solid rgba(0,0,0,0.15);border-radius:6px;'
        f'padding:4px 12px;font-weight:500;">기업 정보 →</a>'
        f'</div>'
        f'<h2 style="font-size:1.6rem;font-weight:800;color:#0a0a18;margin:0 0 6px 0;'
        f'letter-spacing:-0.6px;line-height:1.25;">{corp_name} 협력 적합성 분석 리포트</h2>'
        f'<p style="color:#888;font-size:1.0rem;margin:0 0 20px 0;line-height:1.5;">'
        f'내 기업 프로필과 협력 기업의 요구사항을 비교 분석한 결과입니다.</p>'
        f'<hr style="border:none;border-top:1.5px solid rgba(0,0,0,0.07);margin:0 0 22px 0;">'
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 2 — 3단 비교 헤더 (내 기업 | 종합 적합도 | 협력 기업)
    # ════════════════════════════════════════════════════════════════════════
    emp_str = f"{sr['employees']:,}명" if sr.get("employees") else "-"

    S += (
        f'<div style="display:grid;grid-template-columns:1fr 1fr 1fr;'
        f'border:1.5px solid rgba(0,0,0,0.08);border-radius:14px;overflow:hidden;margin-bottom:24px;">'

        # 왼쪽: 내 기업
        f'<div style="background:rgba(70,90,210,0.04);padding:20px;'
        f'border-right:1.5px solid rgba(0,0,0,0.07);">'
        f'<p style="color:#4060b0;font-size:0.72rem;font-weight:700;letter-spacing:1.2px;'
        f'text-transform:uppercase;margin:0 0 12px 0;">내 기업 프로필</p>'
        f'<div style="display:flex;gap:10px;align-items:flex-start;">'
        f'<div style="width:38px;height:38px;background:rgba(70,90,210,0.1);border-radius:9px;'
        f'display:flex;align-items:center;justify-content:center;font-size:1.15rem;flex-shrink:0;">🏢</div>'
        f'<div style="flex:1;">'
        f'<p style="color:#2a3a80;font-size:1.05rem;font-weight:700;margin:0 0 3px 0;">{p_industry}</p>'
        f'<p style="color:#5060a0;font-size:0.93rem;margin:0 0 8px 0;">{p_category}</p>'
        f'<p style="color:#6a7ab0;font-size:0.9rem;line-height:1.55;margin:0;">{p_desc_short}</p>'
        f'</div></div></div>'

        # 가운데: 종합 적합도
        f'<div style="background:#fff;padding:20px;text-align:center;'
        f'border-right:1.5px solid rgba(0,0,0,0.07);display:flex;flex-direction:column;'
        f'align-items:center;justify-content:center;">'
        f'<p style="color:#888;font-size:0.72rem;font-weight:700;letter-spacing:1.2px;'
        f'text-transform:uppercase;margin:0 0 10px 0;">종합 적합도</p>'
        f'<p style="color:{bar_color};font-size:2.6rem;font-weight:800;margin:0 0 4px 0;'
        f'letter-spacing:-1.5px;line-height:1;">{pct:.0f}%</p>'
        f'<span style="background:{judgment_bg};color:{judgment_color};font-size:0.82rem;'
        f'font-weight:700;border-radius:5px;padding:3px 11px;display:inline-block;margin-bottom:12px;">'
        f'{judgment_label}</span>'
        f'<div style="width:100%;background:rgba(0,0,0,0.08);border-radius:4px;height:7px;overflow:hidden;">'
        f'<div style="background:linear-gradient(90deg,{bar_color},{bar_color}88);'
        f'width:{pct}%;height:100%;border-radius:4px;"></div></div>'
        f'</div>'

        # 오른쪽: 협력 기업
        f'<div style="background:rgba(232,48,90,0.03);padding:20px;">'
        f'<p style="color:#E8305A;font-size:0.72rem;font-weight:700;letter-spacing:1.2px;'
        f'text-transform:uppercase;margin:0 0 12px 0;">협력 기업</p>'
        f'<div style="display:flex;gap:10px;align-items:flex-start;">'
        f'<div style="width:38px;height:38px;background:rgba(232,48,90,0.1);border-radius:9px;'
        f'display:flex;align-items:center;justify-content:center;font-size:1.15rem;flex-shrink:0;">🏭</div>'
        f'<div style="flex:1;">'
        f'<p style="color:#1a0a10;font-size:1.05rem;font-weight:700;margin:0 0 3px 0;">{corp_name}</p>'
        f'<p style="color:#c07080;font-size:0.93rem;margin:0 0 8px 0;">'
        f'{sr.get("market","-")} · 업종 {sr.get("industry","-")}</p>'
        f'<p style="color:#9a6070;font-size:0.9rem;margin:0 0 2px 0;">매출 {rev_str}</p>'
        f'<p style="color:#9a6070;font-size:0.9rem;margin:0;">부채비율 {debt_str}</p>'
        f'</div></div></div>'
        f'</div>'
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 3 — 핵심 분석 요약 (카드 3개)
    # ════════════════════════════════════════════════════════════════════════
    # 판정 카드 body
    if collab_reason:
        card1_body = collab_reason
    elif is_error:
        card1_body = sr.get("reason", "업종 부적합")[:90]
    elif rag_items:
        top_item = max(rag_items, key=lambda x: x[2])
        card1_body = f'{p_category} 솔루션 제공 기업으로서 {corp_name}의 요구사항과 적합성을 보입니다.'
    else:
        card1_body = f'{p_category} 솔루션 적합성 분석 결과입니다.'

    # 핵심 근거 카드 (가장 낮은 RAG 항목의 reason)
    if rag_items:
        key_item = min(rag_items, key=lambda x: x[2])
        card2_body = key_item[3][:90] if key_item[3] else "분석 정보 없음"
    else:
        card2_body = sr.get("reason", "분석 정보 없음")[:90]

    S += (
        f'<p style="color:#1a0a10;font-size:1.15rem;font-weight:700;margin:0 0 12px 0;'
        f'letter-spacing:-0.3px;">핵심 분석 요약</p>'
        f'<div style="display:grid;grid-template-columns:1fr 1fr 1fr;gap:12px;margin-bottom:26px;">'

        # 판정
        f'<div style="background:#fff;border:1.5px solid rgba(0,0,0,0.08);border-radius:12px;padding:18px;">'
        f'<div style="display:flex;align-items:center;gap:8px;margin-bottom:10px;">'
        f'<span style="width:30px;height:30px;background:{judgment_bg};border-radius:7px;'
        f'display:inline-flex;align-items:center;justify-content:center;font-size:1.05rem;'
        f'color:{judgment_color};font-weight:800;">{"✓" if not is_error and total_display>=70 else ("△" if not is_error else "✗")}</span>'
        f'<span style="color:#aaa;font-size:0.78rem;font-weight:600;text-transform:uppercase;letter-spacing:0.8px;">판정</span>'
        f'</div>'
        f'<p style="color:{judgment_color};font-size:1.08rem;font-weight:700;margin:0 0 7px 0;">{judgment_label}</p>'
        f'<p style="color:#5a4050;font-size:0.97rem;line-height:1.6;margin:0;">{card1_body}</p>'
        f'</div>'

        # 핵심 근거
        f'<div style="background:#fff;border:1.5px solid rgba(0,0,0,0.08);border-radius:12px;padding:18px;">'
        f'<div style="display:flex;align-items:center;gap:8px;margin-bottom:10px;">'
        f'<span style="width:30px;height:30px;background:rgba(245,158,11,0.1);border-radius:7px;'
        f'display:inline-flex;align-items:center;justify-content:center;font-size:1.05rem;color:#d97706;">⚠</span>'
        f'<span style="color:#aaa;font-size:0.78rem;font-weight:600;text-transform:uppercase;letter-spacing:0.8px;">핵심 근거</span>'
        f'</div>'
        f'<p style="color:#1a0a10;font-size:1.08rem;font-weight:700;margin:0 0 7px 0;">주요 판단 기준</p>'
        f'<p style="color:#5a4050;font-size:0.97rem;line-height:1.6;margin:0;">{card2_body}</p>'
        f'</div>'

        # 재무 요약
        f'<div style="background:#fff;border:1.5px solid rgba(0,0,0,0.08);border-radius:12px;padding:18px;">'
        f'<div style="display:flex;align-items:center;gap:8px;margin-bottom:10px;">'
        f'<span style="width:30px;height:30px;background:rgba(16,185,129,0.1);border-radius:7px;'
        f'display:inline-flex;align-items:center;justify-content:center;font-size:1.05rem;color:#059669;">↗</span>'
        f'<span style="color:#aaa;font-size:0.78rem;font-weight:600;text-transform:uppercase;letter-spacing:0.8px;">재무 요약</span>'
        f'</div>'
        f'<p style="color:#1a0a10;font-size:1.08rem;font-weight:700;margin:0 0 7px 0;">재무 현황</p>'
        f'<p style="color:#5a4050;font-size:0.97rem;line-height:1.6;margin:0;">'
        f'매출 {rev_str}, 부채비율 {debt_str} — {debt_interp}</p>'
        f'</div>'

        f'</div>'
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 4 — 기업 정보 (주력 사업 | 고객사·파트너 | 협력 신호)
    # ════════════════════════════════════════════════════════════════════════
    if main_products or key_partners or collab_signal:
        cs_map = {
            "positive": ("#1a7a40", "rgba(30,140,70,0.09)", "● 협력 긍정적"),
            "cautious":  ("#b06000", "rgba(180,100,0,0.09)", "● 조건부 협력"),
            "negative":  ("#C4183C", "rgba(200,30,60,0.09)", "● 협력 부정적"),
        }
        cs_color, cs_bg, cs_label = cs_map.get(collab_signal, ("#888", "rgba(0,0,0,0.05)", "● 판단 보류"))
        S += (
            f'<p style="color:#1a0a10;font-size:1.15rem;font-weight:700;margin:0 0 12px 0;'
            f'letter-spacing:-0.3px;">기업 정보</p>'
            f'<div style="display:grid;grid-template-columns:1fr 1fr 1fr;gap:12px;margin-bottom:26px;">'

            f'<div style="background:#fff;border:1.5px solid rgba(0,0,0,0.08);border-radius:12px;padding:18px;">'
            f'<p style="color:#aaa;font-size:0.75rem;font-weight:700;text-transform:uppercase;'
            f'letter-spacing:0.9px;margin:0 0 8px 0;">주력 사업</p>'
            f'<p style="color:#1a0a10;font-size:1.0rem;line-height:1.65;margin:0;">'
            f'{main_products if main_products else "정보 없음"}</p>'
            f'</div>'

            f'<div style="background:#fff;border:1.5px solid rgba(0,0,0,0.08);border-radius:12px;padding:18px;">'
            f'<p style="color:#aaa;font-size:0.75rem;font-weight:700;text-transform:uppercase;'
            f'letter-spacing:0.9px;margin:0 0 8px 0;">주요 고객사·파트너</p>'
            f'<p style="color:#1a0a10;font-size:1.0rem;line-height:1.65;margin:0;">'
            f'{key_partners if key_partners else "공시에서 확인 불가"}</p>'
            f'</div>'

            f'<div style="background:#fff;border:1.5px solid rgba(0,0,0,0.08);border-radius:12px;padding:18px;">'
            f'<p style="color:#aaa;font-size:0.75rem;font-weight:700;text-transform:uppercase;'
            f'letter-spacing:0.9px;margin:0 0 8px 0;">협력 신호</p>'
            f'<span style="background:{cs_bg};color:{cs_color};font-size:0.88rem;font-weight:700;'
            f'border-radius:5px;padding:3px 10px;display:inline-block;margin-bottom:8px;">{cs_label}</span>'
            f'<p style="color:#5a4050;font-size:0.95rem;line-height:1.6;margin:0;">{collab_reason}</p>'
            f'</div>'

            f'</div>'
        )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 5 — 주요 체크 포인트 (RAG 5항목 테이블)
    # ════════════════════════════════════════════════════════════════════════
    _icons = {
        "problem_fit":          ("🔍", "rgba(59,130,246,0.08)"),
        "digital_willingness":  ("💻", "rgba(99,102,241,0.08)"),
        "investment_direction": ("🔗", "rgba(245,158,11,0.08)"),
        "external_adoption":    ("⭐", "rgba(16,185,129,0.08)"),
        "decision_structure":   ("👥", "rgba(239,68,68,0.08)"),
    }
    rag_rows = ""
    for key, label, sc, rsn in rag_items:
        icon, icon_bg = _icons.get(key, ("·", "rgba(0,0,0,0.05)"))
        if sc >= 7:   rel_label, rel_c, rel_bg = "높음", "#1a7a40", "rgba(30,160,80,0.1)"
        elif sc >= 4: rel_label, rel_c, rel_bg = "중간", "#b06000", "rgba(245,158,11,0.1)"
        else:         rel_label, rel_c, rel_bg = "낮음", "#C4183C", "rgba(232,48,90,0.1)"
        bar_w = sc / 10 * 100
        rag_rows += (
            f'<tr style="border-bottom:1.5px solid rgba(0,0,0,0.05);">'
            f'<td style="padding:14px 16px;vertical-align:middle;">'
            f'<div style="display:flex;align-items:flex-start;gap:12px;">'
            f'<span style="width:34px;height:34px;background:{icon_bg};border-radius:9px;'
            f'display:inline-flex;align-items:center;justify-content:center;'
            f'font-size:1.05rem;flex-shrink:0;">{icon}</span>'
            f'<div>'
            f'<p style="color:#1a0a10;font-size:1.02rem;font-weight:600;margin:0 0 3px 0;">{label}</p>'
            f'<p style="color:#7a5060;font-size:0.97rem;margin:0;line-height:1.55;">{rsn}</p>'
            f'</div></div></td>'
            f'<td style="padding:14px 16px;text-align:right;vertical-align:middle;white-space:nowrap;">'
            f'<div style="display:flex;align-items:center;gap:8px;justify-content:flex-end;">'
            f'<span style="color:#bbb;font-size:0.9rem;">{sc}/10</span>'
            f'<div style="width:64px;background:rgba(0,0,0,0.08);border-radius:3px;height:5px;">'
            f'<div style="background:linear-gradient(90deg,#E8305A,#FF8080);'
            f'width:{bar_w}%;height:100%;border-radius:3px;"></div></div>'
            f'<span style="background:{rel_bg};color:{rel_c};font-size:0.8rem;font-weight:700;'
            f'border-radius:5px;padding:2px 9px;min-width:38px;text-align:center;'
            f'display:inline-block;">{rel_label}</span>'
            f'</div></td></tr>'
        )

    S += (
        f'<p style="color:#1a0a10;font-size:1.15rem;font-weight:700;margin:0 0 12px 0;'
        f'letter-spacing:-0.3px;">주요 체크 포인트</p>'
        f'<div style="background:#fff;border:1.5px solid rgba(0,0,0,0.08);'
        f'border-radius:12px;overflow:hidden;margin-bottom:26px;">'
        f'<table style="width:100%;border-collapse:collapse;">'
        f'<thead><tr style="background:rgba(0,0,0,0.025);'
        f'border-bottom:1.5px solid rgba(0,0,0,0.08);">'
        f'<th style="text-align:left;padding:11px 16px;color:#aaa;font-size:0.75rem;'
        f'font-weight:700;text-transform:uppercase;letter-spacing:0.9px;">항목 / 분석 근거</th>'
        f'<th style="text-align:right;padding:11px 16px;color:#aaa;font-size:0.75rem;'
        f'font-weight:700;text-transform:uppercase;letter-spacing:0.9px;">관련성</th>'
        f'</tr></thead>'
        f'<tbody>{rag_rows}</tbody>'
        f'</table></div>'
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 6 — 다음 단계 제안 / 재검토 조건
    # ════════════════════════════════════════════════════════════════════════
    if is_error:
        section_title = "이런 변화가 있다면 재검토 가능합니다"
        _retry_map = {
            "problem_fit":          ("🔍", "문제 명시 확인",   "사업보고서에서 솔루션이 해결할 수 있는 운영 문제가 명시된다면"),
            "digital_willingness":  ("💻", "디지털 투자 공시", "IT·디지털 전환 투자 또는 스마트 시스템 도입 공시가 추가된다면"),
            "investment_direction": ("🔗", "신사업 방향 일치", "신사업 투자 방향이 솔루션 카테고리와 겹치는 사업 계획이 발표된다면"),
            "external_adoption":    ("⭐", "외부 도입 이력",   "외부 소프트웨어·컨설팅 도입 이력이 공시에서 확인된다면"),
            "decision_structure":   ("👥", "의사결정 구조",    "구매 담당 조직 또는 분권화된 의사결정 구조가 확인된다면"),
        }
        action_items = [(icon, t, d) for key, (icon, t, d) in _retry_map.items()
                        if key in rag_d and isinstance(rag_d[key], dict) and rag_d[key].get("score", 10) < 5]
        if not action_items:
            action_items = [("💡", "추가 공시 확인", "향후 공시에서 디지털 전환 또는 외부 솔루션 도입 이력이 확인된다면")]
    else:
        section_title = "다음 단계 제안"
        action_items = [
            ("📋", "사업보고서·공시 심층 검토",  f"최신 사업보고서와 투자 관련 공시를 분석해 {corp_name}의 구체적 니즈를 확인하세요."),
            ("🤝", "접점 발굴",                  f"구매/IT 담당 조직 및 담당자 정보를 확보하세요."),
            ("📄", "맞춤형 제안서 준비",          f"{corp_name}의 니즈에 맞춘 솔루션 제안서를 작성하세요."),
            ("📧", "후속 커뮤니케이션",           "이메일 또는 LinkedIn을 통해 연결을 시도하세요."),
        ]
    steps_html = ""
    for i, (icon, title, desc) in enumerate(action_items, 1):
        steps_html += (
            f'<div style="display:flex;gap:14px;align-items:flex-start;'
            f'padding:14px 0;border-bottom:1.5px solid rgba(0,0,0,0.05);">'
            f'<span style="width:30px;height:30px;background:#E8305A;border-radius:50%;'
            f'display:inline-flex;align-items:center;justify-content:center;'
            f'color:white;font-size:0.88rem;font-weight:700;flex-shrink:0;">{i}</span>'
            f'<div>'
            f'<p style="color:#1a0a10;font-size:1.02rem;font-weight:700;margin:0 0 3px 0;">{title}</p>'
            f'<p style="color:#7a5060;font-size:0.97rem;margin:0;line-height:1.55;">{desc}</p>'
            f'</div></div>'
        )

    S += (
        f'<p style="color:#1a0a10;font-size:1.15rem;font-weight:700;margin:0 0 12px 0;'
        f'letter-spacing:-0.3px;">{section_title}</p>'
        f'<div style="background:#fff;border:1.5px solid rgba(0,0,0,0.08);'
        f'border-radius:12px;padding:4px 18px;margin-bottom:26px;">'
        f'{steps_html}</div>'
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 7 — 푸터
    # ════════════════════════════════════════════════════════════════════════
    _source_label_map = {
        "DB":         ("데이터 출처: DART 사업보고서 (로컬 DB)", "#ccc"),
        "DART 실시간": ("데이터 출처: DART 실시간 조회", "#6080c0"),
        "수동 입력":   ("데이터 출처: 사용자 직접 입력 (재무 데이터 미검증)", "#b06000"),
    }
    _src_text, _src_color = _source_label_map.get(_src_tag, ("데이터 출처: DART", "#ccc"))
    if not is_error:
        rank_info = f" · {rank_badge.strip()} · DB {total_db:,}개 중 {pct_label}" if rank_badge else f" · DB {total_db:,}개 중 {pct_label}"
    else:
        rank_info = ""
    S += (
        f'<div style="display:flex;flex-wrap:wrap;gap:12px;align-items:center;'
        f'padding-top:14px;border-top:1.5px solid rgba(0,0,0,0.07);">'
        f'<span style="color:#ccc;font-size:0.83rem;">분석일: {today}</span>'
        f'<span style="color:#ddd;">·</span>'
        f'<span style="color:{_src_color};font-size:0.83rem;">{_src_text}</span>'
        f'{rank_info}'
        f'</div>'
    )

    S += '</div>'
    return S


# ── 대시보드 탭 ───────────────────────────────────────────────────────────────
with tab_dashboard:
    _last_ind  = st.session_state.last_industry
    _cat_opts  = INDUSTRY_TO_CATEGORIES.get(_last_ind, list(WEIGHT_PROFILES.keys()))
    _ind_keys  = list(INDUSTRY_MAP.keys())
    _ind_idx   = _ind_keys.index(_last_ind) if _last_ind in _ind_keys else 0
    _cat_idx   = _cat_opts.index(st.session_state.last_category) if st.session_state.last_category in _cat_opts else 0

    # ── 내 기업 프로필 등록 ────────────────────────────────────────────────────
    with st.form("profile_form"):
        _saved_badge = (
            '<span style="display:flex;align-items:center;gap:5px;'
            'color:#1a7a40;font-size:0.85rem;font-weight:700;">'
            '<svg width="16" height="16" viewBox="0 0 16 16" fill="none">'
            '<circle cx="8" cy="8" r="7.5" stroke="#1a7a40" stroke-width="1.2"/>'
            '<polyline points="4.5,8.5 7,11 11.5,5.5" stroke="#1a7a40" stroke-width="1.5"'
            ' stroke-linecap="round" stroke-linejoin="round"/></svg>'
            '저장됨</span>'
            if st.session_state.profile_saved else ''
        )
        st.markdown(
            f'<div style="display:flex;justify-content:space-between;align-items:center;'
            f'margin:4px 0 10px 2px;">'
            f'<p style="color:rgba(255,255,255,0.65);font-size:0.82rem;font-weight:700;'
            f'letter-spacing:1.3px;text-transform:uppercase;margin:0;">내 기업 프로필</p>'
            f'{_saved_badge}</div>',
            unsafe_allow_html=True,
        )
        pc1, pc2 = st.columns([2, 4.6])
        with pc1:
            st.markdown(
                '<p style="color:#3a1020;font-size:0.95rem;font-weight:600;margin:0 0 4px 0;">'
                '타깃 산업군 <span style="color:#E8305A;font-weight:800;">*</span></p>',
                unsafe_allow_html=True,
            )
            p_industry = st.selectbox("타깃 산업군", _ind_keys, index=_ind_idx,
                                       label_visibility="collapsed")
            st.markdown(
                '<p style="color:#3a1020;font-size:0.95rem;font-weight:600;margin:6px 0 4px 0;">'
                '제품 카테고리 <span style="color:#E8305A;font-weight:800;">*</span></p>',
                unsafe_allow_html=True,
            )
            p_category = st.selectbox("제품 카테고리", _cat_opts, index=_cat_idx,
                                       label_visibility="collapsed")
        with pc2:
            st.markdown(
                '<p style="color:#3a1020;font-size:0.95rem;font-weight:600;margin:0 0 4px 0;">'
                '주요 제공 가치 / 제품 설명 <span style="color:#E8305A;font-weight:800;">*</span></p>',
                unsafe_allow_html=True,
            )
            p_product_desc = st.text_area(
                "주요 제공 가치 / 제품 설명",
                value=st.session_state.last_product_desc,
                height=100,
                placeholder="우리 솔루션이 해결하는 고객의 근본 문제.\n예: 중소기업 HR 자동화 SaaS로 채용·급여 통합 관리, 인사팀 업무 70% 절감.",
                label_visibility="collapsed",
            )
            _sp, _bc = st.columns([3, 1])
            with _bc:
                profile_btn = st.form_submit_button("프로필 등록", use_container_width=True)

    # ── 필터 탐색 ──────────────────────────────────────────────────────────────
    with st.form("main_form"):
        st.markdown("""
        <p style="color:rgba(255,255,255,0.65);font-size:0.82rem;font-weight:700;letter-spacing:1.3px;
                  text-transform:uppercase;margin:6px 0 8px 2px;">필터 탐색</p>
        """, unsafe_allow_html=True)
        fc1, fc2, fc3 = st.columns([3, 1, 1])
        with fc1:
            _min_deal_str = st.text_input(
                "최소 거래 규모 (원)",
                value="5,000,000",
                placeholder="예: 5,000,000",
            )
            try:
                min_deal = int(_min_deal_str.replace(",", "").replace(" ", "").replace("원", ""))
                min_deal = max(1_000_000, min(10_000_000_000, min_deal))
            except (ValueError, AttributeError):
                min_deal = 5_000_000
        with fc2:
            top_n = st.slider("상위 N개 표시", 5, 50, 20)
        with fc3:
            st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
            run_btn = st.form_submit_button("분석 시작", use_container_width=True)

    # ── 기업명 직접 검색 ──────────────────────────────────────────────────────
    with st.form("search_form"):
        st.markdown("""
        <p style="color:rgba(255,255,255,0.65);font-size:0.82rem;font-weight:700;letter-spacing:1.3px;
                  text-transform:uppercase;margin:6px 0 8px 2px;">기업명 직접 검색</p>
        """, unsafe_allow_html=True)
        sb1, sb2 = st.columns([4, 1])
        with sb1:
            corp_name_input = st.text_input("기업명", placeholder="예: 카카오, 더존비즈온")
        with sb2:
            st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
            search_btn = st.form_submit_button("기업 분석", use_container_width=True)

    st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)
    _r_now  = st.session_state.stage2_result
    _sr_now = st.session_state.search_result
    if _r_now is not None and len(_r_now) > 0:
        m1, m2, m3, m4 = st.columns(4)
        with m1: st.metric("DB 전체 기업", f"{total_in_db:,}개")
        with m2: st.metric("1단계 통과",   f"{st.session_state.stage1_count:,}개")
        with m3: st.metric("2단계 선별",   f"{len(_r_now)}개")
        with m4: st.metric("평균 종합 점수", f"{_r_now['total_score'].mean():.1f}점")
    elif _sr_now is not None and not _sr_now.get("error"):
        m1, m2, m3, m4 = st.columns(4)
        with m1: st.metric("DB 전체 기업", f"{total_in_db:,}개")
        with m2: st.metric("종합 점수",    f"{_sr_now['total']:.1f}점")
        with m3: st.metric("R&D 강도",     f"{_sr_now['rd']:.1f}/10")
        with m4: st.metric("재무 안정성",   f"{_sr_now['debt']:.1f}/10")

    st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)

    _DB_YEAR = 2025

    # ── 프로필 등록 버튼 핸들러 ────────────────────────────────────────────────
    if profile_btn:
        if not p_product_desc.strip():
            st.warning("제품 설명을 입력해주세요.")
        else:
            st.session_state.last_industry    = p_industry
            st.session_state.last_category    = p_category
            st.session_state.last_product_desc = p_product_desc
            st.session_state.profile_saved    = True
            st.rerun()

    # ── 분석 시작 버튼 핸들러 ─────────────────────────────────────────────────
    if run_btn:
        if not st.session_state.profile_saved:
            st.warning("먼저 '내 기업 프로필'을 작성하고 [프로필 등록]을 눌러주세요.")
        else:
            _industry    = st.session_state.last_industry
            _category    = st.session_state.last_category
            _product_desc = st.session_state.last_product_desc
            with st.spinner("1단계 필터링 중..."):
                df1 = stage1_macro_filter(industry_category=_industry,
                                           min_deal_size=int(min_deal), year=_DB_YEAR)
            st.session_state.stage1_count = len(df1)
            if len(df1) == 0:
                st.warning("1단계 통과 기업이 없습니다. 조건을 완화해보세요.")
            else:
                df_scored = compute_numeric_scores(df1)
                try:
                    weights = get_weights(_category, _product_desc)
                except Exception as e:
                    if "CREDIT_EXHAUSTED" in str(e) or "credit" in str(e).lower():
                        st.error("Anthropic API 크레딧이 부족합니다. [console.anthropic.com](https://console.anthropic.com) 에서 충전 후 재시도하세요.")
                        st.stop()
                    weights = {"rd_score": 0.33, "debt_score": 0.33, "rag_score": 0.34}
                st.session_state.weights_used = weights

                pre_n = min(max(top_n * 6, 60), 160)
                df_scored["_pre"] = df_scored["rd_score"] * 0.5 + df_scored["debt_score"] * 0.5
                df_candidates = df_scored.nlargest(pre_n, "_pre").drop(columns=["_pre"]).reset_index(drop=True)
                rows = df_candidates.to_dict("records")

                def _score_one(row):
                    biz = fetch_business_text(row["corp_code"], year=_DB_YEAR)
                    rd  = row.get("rd_expense")
                    rd_val = int(rd) if (rd is not None and not pd.isna(rd)) else None
                    rag = score_rag_single(
                        row["corp_name"], biz, _product_desc, _category,
                        industry_code=str(row.get("industry", "")),
                        revenue=int(row.get("revenue") or 0),
                        rd_expense=rd_val,
                    )
                    return row["corp_code"], row["corp_name"], rag

                _est = max(30, (len(rows) // 15) * 3)
                _est_str = f"약 {_est}초" if _est < 90 else f"약 {_est // 60}~{_est // 60 + 1}분"
                prog = st.progress(0, text=f"Claude RAG 병렬 분석 중... (0/{len(rows)}) — {_est_str} 소요")
                rag_results, excluded, done = [], [], 0
                _credit_error = False

                with ThreadPoolExecutor(max_workers=15) as executor:
                    futures = {executor.submit(_score_one, r): r for r in rows}
                    for future in as_completed(futures):
                        if _credit_error:
                            future.cancel()
                            continue
                        done += 1
                        try:
                            corp_code, corp_name, rag = future.result()
                        except RuntimeError as e:
                            if "CREDIT_EXHAUSTED" in str(e):
                                _credit_error = True
                                prog.empty()
                                st.error("Anthropic API 크레딧이 부족합니다. [console.anthropic.com](https://console.anthropic.com) 에서 충전 후 재시도하세요.")
                                continue
                            done_name = futures[future].get("corp_name", "?")
                            excluded.append({"name": done_name, "reason": str(e)})
                            prog.progress(done / len(rows), text=f"Claude RAG 병렬 분석 중... ({done}/{len(rows)}) — {_est_str} 소요")
                            continue
                        except Exception as e:
                            done_name = futures[future].get("corp_name", "?")
                            excluded.append({"name": done_name, "reason": str(e)})
                            prog.progress(done / len(rows), text=f"Claude RAG 병렬 분석 중... ({done}/{len(rows)}) — {_est_str} 소요")
                            continue
                        prog.progress(done / len(rows), text=f"({done}/{len(rows)}) {corp_name} — {_est_str} 소요")
                        if not rag.get("is_valid", True):
                            reason = rag.get("invalid_reason", "업종 부적합")
                            excluded.append({"name": corp_name, "reason": reason})
                            continue
                        rag_score = sum(v["score"] for v in rag.values()
                                        if isinstance(v, dict) and "score" in v) / len(RAG_LABEL)
                        rag_results.append({"corp_code": corp_code, "rag_score": round(rag_score, 2),
                                             "rag_detail": json.dumps(rag, ensure_ascii=False)})

                if not _credit_error:
                    prog.empty()
                st.session_state.excluded_count   = len(excluded)
                st.session_state.excluded_details = excluded
                if rag_results:
                    rag_df = pd.DataFrame(rag_results)
                    merged = df_candidates.merge(rag_df, on="corp_code", how="inner")
                    merged["total_score"] = (
                        merged["rd_score"]   * weights["rd_score"] +
                        merged["debt_score"] * weights["debt_score"] +
                        merged["rag_score"]  * weights["rag_score"]
                    ) * 10
                    _n_sectors = merged["industry"].fillna("00").astype(str).str[:2].nunique()
                    _sector_lim = max(4, -(-top_n // max(_n_sectors, 1)))
                    result = run_stage3(merged, am_slots=top_n, top_n=top_n, sector_limit=_sector_lim)
                    _save_stage2_results(result, _product_desc, _category, _DB_YEAR)
                    st.session_state.stage2_result = result
                    st.session_state.search_result = None
                    st.session_state.report_result = None
                if not _credit_error:
                    st.rerun()

    # ── 기업 분석 버튼 핸들러 ─────────────────────────────────────────────────
    def _run_search_analysis(target: pd.Series, product_desc: str, category: str, source_label: str = ""):
        """target Series로 RAG 분석 실행 후 search_result 저장."""
        rag = None
        try:
            with st.spinner(f"{target['corp_name']} 분석 중..."):
                biz = fetch_business_text(target["corp_code"], year=_DB_YEAR)
                rag = score_rag_single(
                    target["corp_name"], biz, product_desc, category,
                    industry_code=str(target.get("industry", "")),
                    revenue=int(target.get("revenue") or 0),
                    rd_expense=int(target["rd_expense"]) if (target.get("rd_expense") and not pd.isna(target["rd_expense"])) else None,
                )
        except RuntimeError as e:
            if "CREDIT_EXHAUSTED" in str(e):
                st.error("Anthropic API 크레딧이 부족합니다. [console.anthropic.com](https://console.anthropic.com) 에서 충전 후 재시도하세요.")
                st.stop()
            st.session_state.search_result = {
                "error": True, "name": target["corp_name"],
                "reason": f"분석 오류: {e}", "source": source_label,
                "revenue": int(target.get("revenue") or 0),
                "debt_ratio": float(target.get("debt_ratio") or 0),
                "rag_detail": "{}",
            }
            return
        except Exception as e:
            st.session_state.search_result = {
                "error": True, "name": target["corp_name"],
                "reason": f"분석 중 오류: {e}", "source": source_label,
                "revenue": int(target.get("revenue") or 0),
                "debt_ratio": float(target.get("debt_ratio") or 0),
                "rag_detail": "{}",
            }
            return
        if rag is None:
            return
        if not rag.get("is_valid", True):
            st.session_state.search_result = {
                "error": True, "name": target["corp_name"],
                "reason": rag.get("invalid_reason", "업종 부적합"), "source": source_label,
                "revenue": int(target.get("revenue") or 0),
                "debt_ratio": float(target.get("debt_ratio") or 0),
                "rag_detail": json.dumps(rag, ensure_ascii=False),
            }
        else:
            rag_score = sum(v["score"] for v in rag.values()
                            if isinstance(v, dict) and "score" in v) / len(RAG_LABEL)
            w2    = WEIGHT_PROFILES.get(category) or {"rd_score": 0.33, "debt_score": 0.33, "rag_score": 0.34}
            rd_s  = min(((target.get("rd_expense") or 0) / (target["revenue"] or 1)) * 500, 10.0)
            dbt_s = max(0.0, 10.0 - (target.get("debt_ratio") or 100) / 20)
            total = (rd_s * w2["rd_score"] + dbt_s * w2["debt_score"] + rag_score * w2["rag_score"]) * 10
            st.session_state.search_result = {
                "error": False, "name": target["corp_name"],
                "total": total, "rd": rd_s, "debt": dbt_s, "rag": rag_score,
                "industry":   str(target.get("industry", "-")),
                "market":     str(target.get("market_type", "-")),
                "revenue":    int(target.get("revenue") or 0),
                "employees":  int(target.get("employees") or 0),
                "debt_ratio": float(target.get("debt_ratio") or 0),
                "corp_code":  str(target["corp_code"]),
                "source":     source_label,
                "rag_detail": json.dumps(rag, ensure_ascii=False),
            }

    if search_btn:
        if not st.session_state.profile_saved:
            st.warning("먼저 '내 기업 프로필'을 작성하고 [프로필 등록]을 눌러주세요.")
        elif not corp_name_input.strip():
            st.warning("기업명을 입력해주세요.")
        else:
            _s_product_desc = st.session_state.last_product_desc
            _s_category     = st.session_state.last_category
            st.session_state.stage2_result    = None
            st.session_state.search_result    = None
            st.session_state.excluded_details = []
            st.session_state.excluded_count   = 0
            st.session_state.weights_used     = {}
            st.session_state.report_result    = None
            st.session_state.show_manual_form = False
            st.session_state.manual_corp_name = corp_name_input.strip()

            # 1) 로컬 DB 검색
            try:
                conn  = get_connection()
                cur   = conn.execute("SELECT * FROM companies WHERE corp_name LIKE ? AND year=?",
                                      (f"%{corp_name_input.strip()}%", _DB_YEAR))
                rows  = cur.fetchall()
                cols  = [d[0] for d in cur.description]
                conn.close()
            except Exception:
                rows, cols = [], []

            if rows:
                target = pd.DataFrame(rows, columns=cols).iloc[0]
                _run_search_analysis(target, _s_product_desc, _s_category, source_label="DB")
                st.rerun()
            else:
                # 2) DART API로 corp_code 조회 → DB 재검색 (한글↔영문명 불일치 해결)
                with st.spinner(f"'{corp_name_input.strip()}' DART에서 법인 코드 검색 중..."):
                    dart_hits = search_corp_by_name(corp_name_input.strip())
                if dart_hits:
                    # 2a) 찾은 corp_code들로 DB 재검색
                    dart_codes = [h["corp_code"] for h in dart_hits]
                    placeholders = ",".join("?" * len(dart_codes))
                    try:
                        conn2 = get_connection()
                        cur2  = conn2.execute(
                            f"SELECT * FROM companies WHERE corp_code IN ({placeholders}) AND year=?",
                            dart_codes + [_DB_YEAR],
                        )
                        rows2 = cur2.fetchall()
                        cols2 = [d[0] for d in cur2.description]
                        conn2.close()
                    except Exception:
                        rows2, cols2 = [], []
                    if rows2:
                        target = pd.DataFrame(rows2, columns=cols2).iloc[0]
                        _run_search_analysis(target, _s_product_desc, _s_category, source_label="DB")
                        st.rerun()
                    else:
                        # 2b) DB에도 없으면 실시간 수집
                        dart_row = dart_hits[0]
                        with st.spinner(f"DART 데이터 실시간 수집 중 — {dart_row['corp_name']}..."):
                            rt = fetch_corp_realtime(dart_row["corp_code"], dart_row["corp_name"], year=_DB_YEAR - 1)
                        if rt:
                            target = pd.Series(rt)
                            _run_search_analysis(target, _s_product_desc, _s_category, source_label="DART 실시간")
                            st.rerun()
                        else:
                            st.session_state.show_manual_form = True
                            st.rerun()
                else:
                    # 3) 수동 입력 폼으로 안내 (DART 공시 없는 기업)
                    st.session_state.show_manual_form = True
                    st.rerun()

    # ── 수동 입력 폼 (DB·DART 모두 없을 때) ───────────────────────────────────
    if st.session_state.show_manual_form:
        _mc = st.session_state.manual_corp_name
        st.markdown(
            f'<div style="background:rgba(232,48,90,0.04);border:1.5px solid rgba(232,48,90,0.18);'
            f'border-radius:14px;padding:16px 20px;margin-bottom:12px;">'
            f'<p style="color:#C4183C;font-weight:700;font-size:1.0rem;margin:0 0 4px 0;">'
            f'<b>"{_mc}"</b> — DART 공시 데이터 없음</p>'
            f'<p style="color:#9a5060;font-size:0.92rem;margin:0;">비상장·소규모 기업이거나 공시 미제출 기업입니다. '
            f'기업 정보를 직접 입력하면 RAG 분석을 진행할 수 있습니다.</p>'
            f'</div>',
            unsafe_allow_html=True,
        )
        with st.form("manual_form"):
            mf1, mf2, mf3 = st.columns(3)
            with mf1:
                m_revenue_억 = st.number_input("매출액 (억 원)", min_value=0, value=500, step=100)
                m_industry   = st.selectbox("업종", list(INDUSTRY_MAP.keys()), index=0)
            with mf2:
                m_employees  = st.number_input("임직원 수 (명)", min_value=0, value=100, step=10)
                m_debt_ratio = st.number_input("부채비율 (%)", min_value=0, value=100, step=10)
            with mf3:
                m_rd_억      = st.number_input("R&D 투자액 (억 원, 없으면 0)", min_value=0, value=0, step=10)
                m_market     = st.selectbox("시장 구분", ["비상장", "KOSPI", "KOSDAQ"])
            m_biz_desc = st.text_area(
                "기업 설명 (사업 내용, 주요 고객, 도입 솔루션 등 알고 있는 정보)",
                height=80,
                placeholder="예: 식품 제조 기업으로 스마트팩토리 도입을 추진 중이며, ERP 시스템을 운영하고 있음.",
            )
            manual_submit = st.form_submit_button("이 정보로 분석 시작", use_container_width=True)

        if manual_submit:
            if not st.session_state.profile_saved:
                st.warning("먼저 프로필을 등록해주세요.")
            else:
                _s_product_desc = st.session_state.last_product_desc
                _s_category     = st.session_state.last_category
                _m_revenue      = int(m_revenue_억 * 1_0000_0000)
                _m_rd           = int(m_rd_억 * 1_0000_0000)
                _ind_codes      = INDUSTRY_MAP.get(m_industry, [])
                _ind_code       = _ind_codes[0] if _ind_codes else "00"
                try:
                    with st.spinner(f"{_mc} 분석 중 (수동 입력)..."):
                        rag = score_rag_single(
                            _mc, m_biz_desc, _s_product_desc, _s_category,
                            industry_code=_ind_code,
                            revenue=_m_revenue,
                            rd_expense=_m_rd or None,
                        )
                except RuntimeError as e:
                    if "CREDIT_EXHAUSTED" in str(e):
                        st.error("Anthropic API 크레딧이 부족합니다.")
                        st.stop()
                    st.error(f"분석 오류: {e}")
                    rag = None
                except Exception as e:
                    st.error(f"분석 오류: {e}")
                    rag = None

                if rag:
                    rag_score = sum(v["score"] for v in rag.values()
                                    if isinstance(v, dict) and "score" in v) / len(RAG_LABEL)
                    w2    = WEIGHT_PROFILES.get(_s_category) or {"rd_score": 0.33, "debt_score": 0.33, "rag_score": 0.34}
                    rd_s  = min((_m_rd / (_m_revenue or 1)) * 500, 10.0)
                    dbt_s = max(0.0, 10.0 - m_debt_ratio / 20)
                    total = (rd_s * w2["rd_score"] + dbt_s * w2["debt_score"] + rag_score * w2["rag_score"]) * 10
                    st.session_state.search_result = {
                        "error": False, "name": _mc,
                        "total": total, "rd": rd_s, "debt": dbt_s, "rag": rag_score,
                        "industry":   _ind_code,
                        "market":     m_market,
                        "revenue":    _m_revenue,
                        "employees":  int(m_employees),
                        "debt_ratio": float(m_debt_ratio),
                        "corp_code":  "",
                        "source":     "수동 입력",
                        "rag_detail": json.dumps(rag, ensure_ascii=False),
                    }
                    st.session_state.show_manual_form = False
                    st.rerun()

    # ── 결과 렌더링 ──
    result_df = st.session_state.stage2_result
    sr        = st.session_state.search_result

    if result_df is not None:
        w           = st.session_state.weights_used
        weight_str  = (f"R&D {w.get('rd_score',0):.0%} · 부채 {w.get('debt_score',0):.0%} · RAG {w.get('rag_score',0):.0%}"
                       if w else "")
        exc_details = st.session_state.excluded_details
        exc         = st.session_state.excluded_count

        if len(result_df) > 0:
            cards_html = "".join(result_card(i + 1, row) for i, (_, row) in enumerate(result_df.iterrows()))
        else:
            cards_html = '<p style="color:#d0a0aa;font-size:1.15rem;padding:24px 0 12px;text-align:center;">결과가 없습니다. 조건을 완화해보세요.</p>'

        st.markdown(
            f'<div style="background:#FFFFFF;border:1px solid rgba(232,48,90,0.1);border-radius:18px;'
            f'padding:20px 22px;box-shadow:0 2px 16px rgba(232,48,90,0.06);">'
            f'<div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:14px;">'
            f'<span style="color:#E8305A;font-size:1.0rem;font-weight:700;letter-spacing:1px;text-transform:uppercase;">필터 탐색 결과</span>'
            f'<div style="text-align:right;">'
            f'<div style="color:#b07080;font-size:1.0rem;">{weight_str}</div>'
            f'<div style="color:#d0a8b0;font-size:0.78rem;margin-top:2px;">카테고리별 가중치 상이 적용</div>'
            f'</div></div>'
            f'{cards_html}</div>',
            unsafe_allow_html=True,
        )

        if exc_details:
            exc_rows_html = "".join(
                f'<tr style="border-bottom:1px solid rgba(232,48,90,0.06);">'
                f'<td style="padding:7px 12px;color:#1a0a10;font-size:0.9rem;font-weight:600;white-space:nowrap;">'
                f'{item.get("name","?") if isinstance(item,dict) else item}</td>'
                f'<td style="padding:7px 12px;">'
                f'<span style="background:rgba(232,48,90,0.07);color:#b04060;border-radius:5px;'
                f'padding:2px 8px;font-size:0.78rem;font-weight:600;white-space:nowrap;">2단계 Claude 분석</span>'
                f'</td>'
                f'<td style="padding:7px 12px;color:#7a4050;font-size:0.88rem;line-height:1.4;">'
                f'{(item.get("reason","") if isinstance(item,dict) else "")}'
                f'</td>'
                f'</tr>'
                for item in exc_details
            )
            st.markdown(
                f'<details style="margin-top:10px;background:rgba(0,0,0,0.015);'
                f'border:1px solid rgba(232,48,90,0.08);border-radius:12px;padding:10px 14px;">'
                f'<summary style="color:#c07080;font-size:0.85rem;font-weight:700;cursor:pointer;'
                f'text-transform:uppercase;letter-spacing:0.8px;list-style:none;outline:none;">'
                f'Why-Not — 탈락 {exc}개 &nbsp;▾</summary>'
                f'<div style="margin-top:10px;overflow-x:auto;">'
                f'<table style="width:100%;border-collapse:collapse;">'
                f'<thead><tr style="border-bottom:2px solid rgba(232,48,90,0.1);">'
                f'<th style="text-align:left;padding:5px 12px;color:#9a4050;font-size:0.78rem;'
                f'text-transform:uppercase;letter-spacing:0.5px;white-space:nowrap;">기업명</th>'
                f'<th style="text-align:left;padding:5px 12px;color:#9a4050;font-size:0.78rem;'
                f'text-transform:uppercase;letter-spacing:0.5px;white-space:nowrap;">탈락 단계</th>'
                f'<th style="text-align:left;padding:5px 12px;color:#9a4050;font-size:0.78rem;'
                f'text-transform:uppercase;letter-spacing:0.5px;">탈락 사유</th>'
                f'</tr></thead>'
                f'<tbody>{exc_rows_html}</tbody>'
                f'</table></div></details>',
                unsafe_allow_html=True,
            )

    elif sr is not None:
        _profile_info = {
            "industry":     st.session_state.last_industry,
            "category":     st.session_state.last_category,
            "product_desc": st.session_state.last_product_desc,
        }
        content = search_panel_html(
            sr,
            profile=_profile_info,
            stage2_df=result_df,
            total_db=total_in_db,
        )
        st.markdown(
            f'<div style="background:#FFFFFF;border:1.5px solid rgba(0,0,0,0.07);border-radius:18px;'
            f'padding:28px 32px;box-shadow:0 2px 20px rgba(0,0,0,0.07);">'
            f'{content}</div>',
            unsafe_allow_html=True,
        )


# ── 보고서 헬퍼 ──────────────────────────────────────────────────────────────
def _strip_md_bold(text: str) -> str:
    text = _re.sub(r'\*\*(.+?)\*\*', r'\1', text, flags=_re.DOTALL)
    text = _re.sub(r'\*([^*\n]+?)\*', r'\1', text)
    return text


def _fill_bold_runs(para, text: str):
    for i, part in enumerate(_re.split(r'\*\*(.+?)\*\*', text)):
        if part:
            para.add_run(part).bold = (i % 2 == 1)


def _make_docx_bytes(text: str, title: str) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading(title, level=0)
    for line in text.split('\n'):
        s = line.strip()
        if s.startswith('## '):
            doc.add_heading(s[3:], level=1)
        elif s.startswith('# '):
            doc.add_heading(s[2:], level=1)
        elif s.startswith('- ') or s.startswith('• '):
            _fill_bold_runs(doc.add_paragraph(style='List Bullet'), s[2:])
        elif s:
            _fill_bold_runs(doc.add_paragraph(), s)
        else:
            doc.add_paragraph()
    buf = _BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _clipboard_btn(text: str):
    text_json = json.dumps(text)
    components.html(
        "<script>"
        "function doCopy(){"
        f"  var t={text_json};"
        "  navigator.clipboard.writeText(t).then(function(){"
        "    var b=document.getElementById('cp-btn');"
        "    b.textContent='✓ 복사됨!';"
        "    b.style.background='linear-gradient(135deg,#1a8a3c,#27ae60)';"
        "    setTimeout(function(){b.textContent='📋 본문 복사';"
        "      b.style.background='linear-gradient(135deg,#C4183C,#E8305A)';},2000);"
        "  });"
        "}"
        "</script>"
        "<button id='cp-btn' onclick='doCopy()' style='"
        "background:linear-gradient(135deg,#C4183C,#E8305A);"
        "color:#fff;border:none;border-radius:10px;"
        "font-weight:600;font-size:1.0rem;"
        "padding:8px 0;width:100%;cursor:pointer;"
        "box-shadow:0 4px 18px rgba(232,48,90,0.35);"
        "font-family:Pretendard,\"Noto Sans KR\",sans-serif;"
        "'>📋 본문 복사</button>",
        height=46,
    )


# ── 보고서 제작 탭 ────────────────────────────────────────────────────────────
with tab_report:
    result_df_r = st.session_state.stage2_result
    sr_r        = st.session_state.search_result

    corp_options  = []
    corp_data_map = {}

    if result_df_r is not None and len(result_df_r) > 0:
        for _, row in result_df_r.iterrows():
            name = row["corp_name"]
            corp_options.append(name)
            corp_data_map[name] = {
                "corp_name":   name,
                "revenue":     int(row.get("revenue") or 0),
                "employees":   int(row.get("employees") or 0),
                "industry":    str(row.get("industry", "-")),
                "market_type": str(row.get("market_type", "-")),
                "total_score": float(row.get("total_score", 0)),
                "rd_score":    float(row.get("rd_score", 0)),
                "debt_score":  float(row.get("debt_score", 0)),
                "rag_score":   float(row.get("rag_score", 0)),
                "rag_detail":  json.loads(row.get("rag_detail") or "{}"),
            }

    if sr_r is not None and not sr_r.get("error"):
        name = sr_r["name"]
        if name not in corp_data_map:
            corp_options.insert(0, name)
            corp_data_map[name] = {
                "corp_name":   name,
                "revenue":     sr_r.get("revenue", 0),
                "employees":   sr_r.get("employees", 0),
                "industry":    sr_r.get("industry", "-"),
                "market_type": sr_r.get("market", "-"),
                "total_score": sr_r.get("total", 0),
                "rd_score":    sr_r.get("rd", 0),
                "debt_score":  sr_r.get("debt", 0),
                "rag_score":   sr_r.get("rag", 0),
                "rag_detail":  json.loads(sr_r.get("rag_detail") or "{}"),
            }

    for _k, _v in [("rp_tags", []), ("rp_generating", False)]:
        if _k not in st.session_state:
            st.session_state[_k] = _v

    if not corp_options:
        st.markdown(
            '<div style="background:#FFFFFF;border:1px solid rgba(232,48,90,0.1);border-radius:18px;'
            'padding:48px;text-align:center;box-shadow:0 2px 16px rgba(232,48,90,0.06);margin-top:8px;">'
            '<p style="color:#E8305A;font-size:1.2rem;font-weight:600;margin:0 0 8px 0;">분석 결과가 없습니다</p>'
            '<p style="color:#c0a0a8;font-size:1.0rem;margin:0;">'
            '대시보드에서 필터 탐색 또는 기업 단독 분석을 먼저 진행해주세요.'
            '</p></div>',
            unsafe_allow_html=True,
        )
    else:
        gen_btn   = False
        regen_btn = False

        rp_left, rp_right = st.columns([3, 7], gap="medium")

        # ── 왼쪽 입력 패널 ───────────────────────────────────────────────────
        with rp_left:
            st.markdown('<div id="rp-left-anchor"></div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="rp-panel-header">'
                '<span style="font-size:1.1rem;">👤</span> 기본 정보'
                '</div>',
                unsafe_allow_html=True,
            )

            st.markdown('<span class="rp-label">제안 유형</span>', unsafe_allow_html=True)
            doc_type_label = st.selectbox(
                "제안 유형", list(DOC_TYPES.values()),
                label_visibility="collapsed", key="rp_doc_type",
            )

            st.markdown('<span class="rp-label">제안 대상 기업</span>', unsafe_allow_html=True)
            selected_corp = st.selectbox(
                "타깃 기업", corp_options,
                label_visibility="collapsed", key="rp_corp",
            )

            st.markdown('<span class="rp-label">핵심 제안 내용</span>', unsafe_allow_html=True)
            product_desc_input = st.text_area(
                "핵심 제안 내용",
                value=st.session_state.get("last_product_desc", ""),
                height=95, max_chars=500,
                label_visibility="collapsed",
                placeholder="우리 제품·서비스의 핵심 가치를 입력하세요",
                key="rp_product_desc",
            )

            st.markdown(
                '<div style="display:flex;align-items:center;margin-top:10px;">'
                '<span class="rp-label" style="margin:0;">☆ 강조 포인트 (선택)</span>'
                '</div>',
                unsafe_allow_html=True,
            )
            if st.session_state.rp_tags:
                chips = "".join(
                    f'<span class="rp-tag">✓ {t}</span>'
                    for t in st.session_state.rp_tags
                )
                st.markdown(f'<div class="rp-tags-row">{chips}</div>', unsafe_allow_html=True)
                if st.button("전체 초기화", key="rp_tag_clear"):
                    st.session_state.rp_tags = []
                    st.rerun()

            tag_c1, tag_c2 = st.columns([5, 2])
            with tag_c1:
                st.text_input(
                    "새 태그", key="rp_new_tag",
                    label_visibility="collapsed",
                    placeholder="강조 포인트 추가...",
                )
            with tag_c2:
                if st.button("+ 추가", key="rp_tag_add", use_container_width=True):
                    t = (st.session_state.get("rp_new_tag") or "").strip()
                    if t and t not in st.session_state.rp_tags:
                        st.session_state.rp_tags.append(t)
                        st.session_state.rp_new_tag = ""
                        st.rerun()

            st.markdown('<span class="rp-label">수신 담당자</span>', unsafe_allow_html=True)
            recipient_role = st.selectbox(
                "수신 담당자", RECIPIENT_ROLES,
                label_visibility="collapsed", key="rp_recipient",
            )

            st.markdown('<div style="height:8px;"></div>', unsafe_allow_html=True)
            gen_btn = st.button(
                "✦ 내용 생성하기",
                use_container_width=True, key="rp_gen_btn",
            )

        # ── 오른쪽 결과 패널 ─────────────────────────────────────────────────
        with rp_right:
            st.markdown('<div id="rp-right-anchor"></div>', unsafe_allow_html=True)
            rp_generating  = st.session_state.get("rp_generating", False)
            report_result  = st.session_state.get("report_result")
            report_raw     = st.session_state.get("report_result_raw", "")
            doc_type_shown = st.session_state.get("report_doc_type", "")
            corp_shown     = st.session_state.get("report_corp_name", "")
            regen_btn      = False

            if rp_generating:
                pending = st.session_state.get("rp_pending_params", {})
                st.markdown(
                    '<div class="rp-right-empty">'
                    '<div class="rp-right-empty-icon">⏳</div>'
                    '<p class="rp-right-empty-title">문서를 작성하고 있습니다</p>'
                    '<p class="rp-right-empty-sub">Claude가 기업 데이터를 분석해<br>맞춤 문서를 생성 중입니다...</p>'
                    '</div>',
                    unsafe_allow_html=True,
                )
                if pending:
                    cd  = pending["corp_data"]
                    dtk = pending["doc_type_key"]
                    try:
                        with st.spinner(f"{pending['doc_type_lbl']} 작성 중..."):
                            raw_text = generate_report(
                                corp_name=cd["corp_name"],
                                revenue=cd["revenue"],
                                employees=cd["employees"],
                                industry=cd["industry"],
                                market_type=cd["market_type"],
                                total_score=cd["total_score"],
                                rd_score=cd["rd_score"],
                                debt_score=cd["debt_score"],
                                rag_score=cd["rag_score"],
                                rag_detail=cd["rag_detail"],
                                product_description=pending["product_desc"],
                                product_category=pending["product_cat"],
                                doc_type=dtk,
                                recipient_role=pending["role"],
                                emphasis=pending["emphasis_str"],
                            )
                        display_text = _strip_md_bold(raw_text)
                        st.session_state.report_result     = display_text
                        st.session_state.report_result_raw = raw_text
                        st.session_state.report_corp_name  = pending["corp_sel"]
                        st.session_state.report_doc_type   = pending["doc_type_lbl"]
                        st.session_state.rp_last_params    = pending
                    except RuntimeError as e:
                        if "CREDIT_EXHAUSTED" in str(e):
                            st.error("Anthropic API 크레딧이 부족합니다. console.anthropic.com에서 충전 후 재시도하세요.")
                        else:
                            st.error(f"오류: {e}")
                    finally:
                        st.session_state.rp_generating = False
                        st.session_state.pop("rp_pending_params", None)
                    st.rerun()

            elif not report_result:
                st.markdown(
                    '<div class="rp-right-empty">'
                    '<div class="rp-right-empty-icon">📄</div>'
                    '<p class="rp-right-empty-title">아직 생성된 문서가 없습니다</p>'
                    '<p class="rp-right-empty-sub">왼쪽에서 기업과 문서 유형을 선택한 후<br>'
                    "'내용 생성하기'를 클릭하세요</p>"
                    '</div>',
                    unsafe_allow_html=True,
                )
            else:
                tabs_html = '<div class="rp-doc-tabs">'
                for dt_label in DOC_TYPES.values():
                    cls = "rp-doc-tab active" if dt_label == doc_type_shown else "rp-doc-tab"
                    tabs_html += f'<span class="{cls}">{dt_label}</span>'
                tabs_html += '</div>'
                st.markdown(tabs_html, unsafe_allow_html=True)

                st.markdown(
                    f'<p class="rp-corp-label">{corp_shown} · Claude 생성 초안</p>',
                    unsafe_allow_html=True,
                )
                st.text_area(
                    "생성 결과", value=report_result,
                    height=385, label_visibility="collapsed",
                )
                st.markdown(
                    '<div class="rp-action-bar">'
                    '<span class="rp-action-hint">AI가 생성한 초안입니다. 내용은 상황에 맞게 수정하여 사용하세요.</span>'
                    '</div>',
                    unsafe_allow_html=True,
                )
                act1, act2 = st.columns([1, 1])
                with act1:
                    regen_btn = st.button(
                        "↺ 다시 생성", key="rp_regenerate",
                        use_container_width=True,
                    )
                with act2:
                    if doc_type_shown == "콜드 메일":
                        _clipboard_btn(report_result)
                    else:
                        docx_bytes = _make_docx_bytes(
                            report_raw or report_result,
                            f"{corp_shown} — {doc_type_shown}",
                        )
                        st.download_button(
                            "⬇ 다운로드 (.docx)",
                            data=docx_bytes,
                            file_name=f"{corp_shown}_{doc_type_shown}.docx",
                            mime="application/vnd.openxmlformats-officedocument"
                                 ".wordprocessingml.document",
                            use_container_width=True,
                            key="rp_download",
                        )

        # ── 생성 트리거 처리 ─────────────────────────────────────────────────
        if gen_btn:
            doc_type_key     = next(k for k, v in DOC_TYPES.items() if v == doc_type_label)
            corp_data        = corp_data_map[selected_corp]
            product_desc_use = product_desc_input or st.session_state.get("last_product_desc", "")
            product_cat_use  = st.session_state.get("last_category", "기타")
            emphasis_str     = ", ".join(st.session_state.rp_tags)

            if not product_desc_use:
                st.warning("핵심 제안 내용을 입력하거나, 대시보드에서 제품 설명을 먼저 입력해주세요.")
            else:
                st.session_state.rp_pending_params = dict(
                    corp_data=corp_data, doc_type_key=doc_type_key,
                    doc_type_lbl=doc_type_label, corp_sel=selected_corp,
                    product_desc=product_desc_use, product_cat=product_cat_use,
                    role=recipient_role, emphasis_str=emphasis_str,
                )
                st.session_state.rp_generating = True
                st.rerun()

        if regen_btn and st.session_state.get("rp_last_params"):
            st.session_state.rp_pending_params = st.session_state.rp_last_params
            st.session_state.rp_generating = True
            st.rerun()


# ── 히스토리 탭 ───────────────────────────────────────────────────────────────
with tab_history:
    st.markdown("""
    <div style="background:#FFFFFF;border:1px solid rgba(232,48,90,0.1);border-radius:18px;
                padding:48px;text-align:center;box-shadow:0 2px 16px rgba(232,48,90,0.06);margin-top:8px;">
        <p style="color:#c07080;font-size:1.2rem;font-weight:600;margin:0 0 8px 0;">준비 중</p>
        <p style="color:#d0b0b8;font-size:1.0rem;margin:0;">
            분석 이력 저장 및 조회 기능은 추후 업데이트 예정입니다.
        </p>
    </div>
    """, unsafe_allow_html=True)


st.markdown("</div>", unsafe_allow_html=True)
